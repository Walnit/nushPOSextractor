import docx, csv

document = docx.Document('table.docx')

headers = ["".join(cell.text.split("\n")).strip() for cell in document.tables[0].rows[0].cells]
data = []

real_headers = ["department", "level", "sem", "code", "type", "title", "description", "mcs", "prerequisites", "preclusions", "corequisites", "hrs", "remarks"]
print("File Headers:", headers)
print("Substitute Headers: ", real_headers)


for prog, tab in enumerate(document.tables):
    print("Progress: " + str(prog+1) + "/" + str(len(document.tables)))
    for row in tab.rows:
        row_data = []
        for cell in row.cells:
            row_data.append("".join(cell.text.split("\n")).strip())

        if row_data != headers:

            if row_data[2] != "": row_data.insert(0, row_data[2][0:2])
            else: row_data.insert(0, "")

            if row_data[0] != "":
                data.append(row_data)
            else:
                for i in range(len(row_data)):
                    data[-1][i] += row_data[i]

with open("pos.csv", 'w', encoding="UTF8", newline='') as f:
    writer = csv.writer(f)

    writer.writerow(real_headers)
    writer.writerows(data)

    f.close()

print("Done!")
